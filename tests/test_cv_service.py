"""Tests for CV parsing and document generation."""

import io

import pytest
from docx import Document

from exceptions import FileProcessingError
from services.cv_service import CVService

SAMPLE_MARKDOWN = """# Jörg Müller
jorg@example.com | +49 123 456

## Professional Summary
Engineer with **10 years** of experience.

## Professional Experience
### Senior Engineer @ Acme
- Improved latency by 20%
- Led a team of 5
"""


def test_parse_txt():
    content = CVService.parse_cv_file("Hello CV ünïcode".encode("utf-8"), "cv.txt")
    assert content == "Hello CV ünïcode"


def test_parse_docx():
    document = Document()
    document.add_paragraph("John Doe")
    document.add_paragraph("Senior Engineer at Acme")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "Kubernetes"
    buffer = io.BytesIO()
    document.save(buffer)

    content = CVService.parse_cv_file(buffer.getvalue(), "cv.docx")
    assert "John Doe" in content
    assert "Senior Engineer at Acme" in content
    assert "Python | Kubernetes" in content


def test_parse_unsupported_extension_raises():
    with pytest.raises(FileProcessingError):
        CVService.parse_cv_file(b"data", "cv.rtf")


def test_parse_corrupt_pdf_raises():
    with pytest.raises(FileProcessingError):
        CVService.parse_cv_file(b"not a pdf", "cv.pdf")


def test_clean_markdown_code_blocks():
    assert CVService.clean_markdown_code_blocks("```markdown\n# Hi\n```") == "# Hi"
    assert CVService.clean_markdown_code_blocks("```\n# Hi\n```") == "# Hi"
    assert CVService.clean_markdown_code_blocks("# Hi") == "# Hi"


def test_generate_pdf_returns_bytes_with_unicode():
    pdf_bytes = CVService.generate_pdf(SAMPLE_MARKDOWN)
    assert pdf_bytes is not None
    assert pdf_bytes.startswith(b"%PDF")


def test_generate_docx_returns_valid_document():
    docx_bytes = CVService.generate_docx(SAMPLE_MARKDOWN)
    assert docx_bytes is not None

    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Jörg Müller" in text
    assert "Improved latency by 20%" in text
