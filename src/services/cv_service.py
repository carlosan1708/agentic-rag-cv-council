import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from pypdf import PdfReader

from exceptions import FileProcessingError
from logger import logger

# PDF Layout Constants
A4_WIDTH_MM = 210
MARGIN_MM = 20
EFFECTIVE_WIDTH = A4_WIDTH_MM - (2 * MARGIN_MM) - 5  # 5mm safety buffer

# Candidate paths for a Unicode-capable TTF font (regular, bold)
UNICODE_FONT_CANDIDATES = [
    (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ),
    (
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
    ),
    (
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ),
]

_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


class CVService:
    @staticmethod
    def parse_cv_file(file_content: bytes, filename: str) -> str:
        """Parses an uploaded CV file (PDF, DOCX or TXT) and returns its text content."""
        try:
            logger.info("Parsing CV file: %s", filename)
            lower_filename = filename.lower()
            if lower_filename.endswith(".pdf"):
                reader = PdfReader(io.BytesIO(file_content))
                content = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
                return content.strip()
            elif lower_filename.endswith(".docx"):
                document = Document(io.BytesIO(file_content))
                parts = [p.text for p in document.paragraphs if p.text.strip()]
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))
                return "\n".join(parts).strip()
            elif lower_filename.endswith(".txt"):
                return file_content.decode("utf-8").strip()
            else:
                raise FileProcessingError(f"Unsupported file format: {filename}. Please upload a PDF, DOCX or TXT file.")
        except FileProcessingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing CV file {filename}: {str(e)}")
            raise FileProcessingError("Failed to read the CV file. Please ensure it is a valid PDF, DOCX or TXT file.") from e

    @staticmethod
    def _sanitize_text_for_pdf(text: str) -> str:
        """Sanitizes text to be compatible with FPDF Latin-1 encoding (fallback path only)."""
        replacements = {
            "–": "-",
            "—": "-",
            "‘": "'",
            "’": "'",
            "“": '"',
            "”": '"',
            "•": "-",
            "…": "...",
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        return text.encode("latin-1", "replace").decode("latin-1")

    @staticmethod
    def clean_markdown_code_blocks(text: str) -> str:
        """Removes markdown code block syntax if present."""
        cleaned = text.strip()
        if cleaned.startswith("```markdown"):
            cleaned = cleaned[11:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        return cleaned.strip()

    @staticmethod
    def _setup_pdf_font(pdf: FPDF) -> str:
        """Registers a Unicode font if available; returns the font family to use."""
        for regular_path, bold_path in UNICODE_FONT_CANDIDATES:
            try:
                pdf.add_font("AppFont", "", regular_path)
                pdf.add_font("AppFont", "B", bold_path)
                return "AppFont"
            except Exception:
                continue
        logger.warning("No Unicode font found; falling back to Helvetica with Latin-1 sanitization.")
        return "helvetica"

    @staticmethod
    def generate_pdf(cv_markdown: str) -> Optional[bytes]:
        """Generates a professional PDF document from Markdown content (A4)."""
        try:
            logger.info("Generating PDF from Markdown...")

            cleaned_cv = CVService.clean_markdown_code_blocks(cv_markdown)

            pdf = FPDF(orientation="P", unit="mm", format="A4")
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_margins(MARGIN_MM, MARGIN_MM, MARGIN_MM)

            font = CVService._setup_pdf_font(pdf)
            if font == "helvetica":
                cleaned_cv = CVService._sanitize_text_for_pdf(cleaned_cv)

            lines = cleaned_cv.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(2)
                    continue

                if line.startswith("# "):
                    # H1 - Name or Main Title
                    pdf.ln(4)
                    pdf.set_x(MARGIN_MM)
                    pdf.set_font(font, "B", 24)
                    pdf.multi_cell(EFFECTIVE_WIDTH, 10, line[2:].upper(), align="C", markdown=True)
                    pdf.ln(6)

                elif line.startswith("## "):
                    # H2 - Section Headers
                    pdf.ln(6)
                    pdf.set_x(MARGIN_MM)
                    pdf.set_font(font, "B", 16)
                    pdf.multi_cell(EFFECTIVE_WIDTH, 8, line[3:].upper(), align="L", markdown=True)

                    # Horizontal Line
                    current_y = pdf.get_y()
                    pdf.set_line_width(0.5)
                    pdf.line(MARGIN_MM, current_y, MARGIN_MM + EFFECTIVE_WIDTH, current_y)
                    pdf.set_line_width(0.2)
                    pdf.ln(4)

                elif line.startswith("### "):
                    # H3 - Subsections
                    pdf.ln(3)
                    pdf.set_x(MARGIN_MM)
                    pdf.set_font(font, "B", 14)
                    pdf.multi_cell(EFFECTIVE_WIDTH, 6, line[3:].strip(), align="L", markdown=True)

                elif line.startswith("- ") or line.startswith("* "):
                    # List Items
                    pdf.set_font(font, size=10)
                    pdf.set_x(MARGIN_MM)
                    current_x = pdf.get_x()
                    pdf.set_x(current_x + 5)
                    pdf.multi_cell(EFFECTIVE_WIDTH - 5, 5, "- " + line[2:], align="L", markdown=True)

                else:
                    # Body Text
                    pdf.set_font(font, size=10)
                    pdf.set_x(MARGIN_MM)
                    pdf.multi_cell(EFFECTIVE_WIDTH, 5, line, align="L", markdown=True)

            logger.info("PDF generated successfully.")
            return bytes(pdf.output())
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return None

    @staticmethod
    def _add_docx_runs(paragraph, text: str) -> None:
        """Adds runs to a docx paragraph, honoring **bold** markdown spans."""
        position = 0
        for match in _BOLD_PATTERN.finditer(text):
            if match.start() > position:
                paragraph.add_run(text[position : match.start()])
            paragraph.add_run(match.group(1)).bold = True
            position = match.end()
        if position < len(text):
            paragraph.add_run(text[position:])

    @staticmethod
    def generate_docx(cv_markdown: str) -> Optional[bytes]:
        """Generates a DOCX document from Markdown content."""
        try:
            logger.info("Generating DOCX from Markdown...")
            cleaned_cv = CVService.clean_markdown_code_blocks(cv_markdown)

            document = Document()
            style = document.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            for line in cleaned_cv.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    document.add_heading(line[2:].strip(), level=0)
                elif line.startswith("## "):
                    document.add_heading(line[3:].strip(), level=1)
                elif line.startswith("### "):
                    document.add_heading(line[4:].strip(), level=2)
                elif line.startswith("- ") or line.startswith("* "):
                    paragraph = document.add_paragraph(style="List Bullet")
                    CVService._add_docx_runs(paragraph, line[2:].strip())
                else:
                    paragraph = document.add_paragraph()
                    CVService._add_docx_runs(paragraph, line)

            buffer = io.BytesIO()
            document.save(buffer)
            logger.info("DOCX generated successfully.")
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return None
